#!/usr/bin/env python3
"""Render a pipeline resume markdown file to a .docx, mirroring render_resume.py's
layout (two-line entries, centered letter-spaced section headers, round bullets)
so the Word file and the PDF match. Same markdown convention -- see
render_resume.py's module docstring for the expected shape.

Usage:
    python render_docx.py <input.md> [-o <output.docx>]
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from render_resume import DATE_RE, split_header, split_italic_header  # noqa: E402

BLUE = RGBColor(0x1F, 0x4E, 0x79)
LINK_BLUE = RGBColor(0x11, 0x55, 0xCC)
FONT = "Times New Roman"


def set_border(paragraph, pos="bottom", sz=8, color="000000", space=2):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{pos}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def add_run(p, text, bold=False, italic=False, size=10.5, color=None, spacing=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rPr.append(sp)
    return r


def add_inline(p, text, base_bold=False, base_italic=False, size=10.5, bold_is_also_italic=False):
    """Minimal **bold**/*italic* inline markdown -> runs."""
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|\*(.+?)\*", text):
        if m.start() > pos:
            add_run(p, text[pos:m.start()], base_bold, base_italic, size)
        if m.group(1) is not None:
            add_run(p, m.group(1), True, base_italic or bold_is_also_italic, size)
        else:
            add_run(p, m.group(2), base_bold, True, size)
        pos = m.end()
    if pos < len(text):
        add_run(p, text[pos:], base_bold, base_italic, size)


def two_col_paragraph(doc, left_text, right_text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.4), alignment=3)  # WD_TAB_ALIGNMENT.RIGHT
    add_inline(p, left_text, base_bold=bold, base_italic=italic, size=size)
    if right_text:
        p.add_run("\t")
        add_run(p, right_text, bold=italic, italic=italic, size=size)
    return p


def start_bullet_paragraph(doc, size=10.5):
    """Create a bullet paragraph with its marker only -- the body text is
    added later via flush_pending(), once every wrapped continuation line
    has been collected. Calling add_inline() per-fragment (an earlier
    version of this fix) breaks any **bold**/*italic* span that straddles
    the manual line-wrap point: the opening ** and closing ** land in
    separate add_inline() calls and neither half matches on its own, so
    both render as literal asterisks instead of bold text."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    add_run(p, "●   ", size=size - 4)
    return p


def build_docx(md_text, out_path):
    lines = [l.strip() for l in md_text.replace("\r\n", "\n").split("\n")
             if l.strip() and not l.strip().startswith("<!--")]
    n = len(lines)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)

    i = 0
    head = []
    while i < n and len(head) < 3 and not lines[i].startswith("#"):
        head.append(lines[i])
        i += 1
    if head:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_run(p, head[0], bold=True, size=20, color=BLUE, spacing=4)
        for line in head[1:3]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, line, size=10)

    first_h2 = True
    section = ""
    # A continuation line (no blank line, no special prefix -- how every
    # draft's bullets and summary paragraphs are manually wrapped for
    # git-diff readability) accumulates as RAW text here and is only run
    # through add_inline() once, on flush -- see start_bullet_paragraph's
    # docstring for why per-fragment inlining breaks bold spans that
    # straddle the wrap point. Blank lines are stripped before this loop
    # even runs, so this pending buffer is the only surviving
    # paragraph-boundary signal.
    pending_p = None
    pending_raw = ""
    pending_kwargs = {}

    def flush():
        nonlocal pending_p, pending_raw, pending_kwargs
        if pending_p is not None:
            add_inline(pending_p, pending_raw, **pending_kwargs)
        pending_p, pending_raw, pending_kwargs = None, "", {}

    while i < n:
        stripped = lines[i]

        if stripped.startswith("## "):
            flush()
            title = stripped[3:].strip()
            key = title.lower()
            section = "skills" if "skill" in key else "other"
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if first_h2:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(3)
                add_run(p, title.upper(), bold=True, size=15, spacing=6)
                first_h2 = False
            else:
                p.paragraph_format.space_before = Pt(9)
                p.paragraph_format.space_after = Pt(4)
                add_run(p, title.upper(), bold=True, size=11, spacing=18)
                set_border(p)
            i += 1
            continue

        if stripped.count("●") >= 2 or stripped.count("•") >= 2:
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(5)
            add_inline(p, stripped, base_bold=True)
            i += 1
            continue

        if stripped.startswith("- ") and section == "skills":
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1.5)
            pending_p, pending_raw = p, stripped[2:]
            pending_kwargs = {"bold_is_also_italic": True}
            i += 1
            continue

        if stripped.startswith("- "):
            flush()
            pending_p = start_bullet_paragraph(doc)
            pending_raw = stripped[2:]
            pending_kwargs = {}
            i += 1
            continue

        if stripped.startswith("**"):
            flush()
            bold, trailing = split_header(stripped)
            next_line = lines[i + 1] if i + 1 < n else ""
            next_is_role_line = next_line.startswith("*") and not next_line.startswith("**")
            bold_trailing_is_date = bool(trailing) and DATE_RE.search(trailing)
            if next_is_role_line and not bold_trailing_is_date:
                role_bold, role_trailing = split_italic_header(next_line)
                two_col_paragraph(doc, bold, trailing, bold=True)
                two_col_paragraph(doc, role_bold, role_trailing, italic=True)
                i += 2
                continue
            if trailing and DATE_RE.search(trailing) and len(trailing) < 40:
                two_col_paragraph(doc, bold, trailing, bold=True)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                add_inline(p, stripped, size=10.5)
            i += 1
            continue

        if stripped.startswith("*") and not stripped.startswith("**"):
            flush()
            italic, trailing = split_italic_header(stripped)
            if trailing and DATE_RE.search(trailing):
                two_col_paragraph(doc, italic, trailing, italic=True)
            else:
                p = doc.add_paragraph()
                pending_p, pending_raw, pending_kwargs = p, stripped, {}
            i += 1
            continue

        # Catch-all: a continuation line with no special prefix. Append
        # its raw text into whatever paragraph is still pending rather
        # than starting a new, wrongly-formatted one.
        if pending_p is not None:
            pending_raw += " " + stripped
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            pending_p, pending_raw, pending_kwargs = p, stripped, {}
        i += 1

    flush()
    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("-o", "--output")
    args = ap.parse_args()

    with open(args.input, encoding="utf-8") as fh:
        md_text = fh.read()

    out = args.output or os.path.splitext(args.input)[0] + ".docx"
    build_docx(md_text, out)
    print(f"DOCX: {os.path.abspath(out)}")


if __name__ == "__main__":
    main()
